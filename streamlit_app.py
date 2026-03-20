# streamlit_app.py
# Minimal UI to exercise both engines using the booklet index from the private repo.

# ── imports (top of file) ──────────────────────────────────────────────────────
import base64
import json
import mimetypes
import os
import requests
import streamlit as st
st.set_page_config(
    page_title="B's Bot",
    page_icon="assets/b2_logo_1024.png",
    layout="wide",
    initial_sidebar_state="collapsed"  # NEW: collapse sidebar by default
)
import time
from typing import Callable, List, Dict, Any
from app.router import route
from mentor.rag.booklet_retriever import extract_signals
# ───────────────────────────────────────────────────────────────────────────────



# === HELPERS ===
# === APP BAR ===
def render_brand_bar_aligned(
    icon_src: str = "assets/b2_logo_1024.png",
    title: str = "B's Bot",
    subhead: str = "Your AI Mentor for EU Capital Markets Law.",
    bar_height_desktop: int = 44,
    bar_height_mobile: int = 38,
    logo_top_nudge_px: int = 0,
    title_nudge_px: int = 0,
    sub_nudge_px: int = 0
) -> None:
    # embed image as data-URI
    try:
        mime, _ = mimetypes.guess_type(icon_src)
        if not mime:
            mime = "image/png"
        with open(icon_src, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")
        img_data_uri = f"data:{mime};base64,{b64}"
    except Exception:
        img_data_uri = icon_src

    st.markdown(
        f"""
<style>
  .lc-appbar {{
    background: #F6F8FC;
    border: 1px solid #E7EAF0;
    border-radius: 10px;
    padding: 8px 10px;
    margin: 6px 0 12px 0;
  }}
  .lc-appbar-row {{
    display: flex;
    align-items: flex-start;             /* lock logo top to text top */
    gap: 10px;
  }}
  .lc-appbar-logo {{
    height: {bar_height_desktop}px;
    width: auto;
    display: block;
    position: relative;
    top: {-logo_top_nudge_px}px;
  }}
  .lc-appbar-text {{
    display: flex;
    flex-direction: column;
    height: {bar_height_desktop}px;      /* SAME height as logo */
    min-height: {bar_height_desktop}px;
  }}
  .lc-appbar-title {{
    margin: 0;
    font-weight: 700;
    font-size: 1.12rem;
    line-height: 1.08;
    position: relative;
    top: {title_nudge_px}px;
    color: #0B1F3B;
  }}
  .lc-appbar-sub {{
    margin: 0;
    font-size: 0.95rem;
    line-height: 1.08;
    margin-top: auto;                    /* pin subtitle to bottom edge */
    position: relative;
    top: {sub_nudge_px}px;
    color: #0B1F3B; opacity: 0.90;
  }}
  @media (max-width: 680px) {{
    .lc-appbar-logo {{ height: {bar_height_mobile}px; }}
    .lc-appbar-text {{ height: {bar_height_mobile}px; min-height: {bar_height_mobile}px; }}
    .lc-appbar-title {{ font-size: 1.02rem; }}
    .lc-appbar-sub {{ font-size: 0.9rem; }}
  }}
</style>

<div class="lc-appbar">
  <div class="lc-appbar-row">
    <img class="lc-appbar-logo" src="{img_data_uri}" alt="B's Bot icon"/>
    <div class="lc-appbar-text">
      <div class="lc-appbar-title">{title}</div>
      <div class="lc-appbar-sub">{subhead}</div>
    </div>
  </div>
</div>
""",
        unsafe_allow_html=True,
    )
# =============================
# Conversation utilities (DRY)
# =============================
def _approx_tokens(s: str) -> int:
    """Rough token estimate (~4 chars / token)."""
    return max(1, len(s) // 4)

def _build_history_preamble(history: List[Dict[str, str]], max_tokens: int = 6000) -> str:
    """
    Create a compact transcript (User/Assistant lines) from the most recent turns
    within a token budget; keeps newest content.
    """
    if not history:
        return ""
    acc, used = [], 0
    for turn in reversed(history):
        line = f'{"User" if turn["role"]=="user" else "Assistant"}: {turn["content"]}'.strip()
        t = _approx_tokens(line)
        if used + t > max_tokens:
            break
        acc.append(line); used += t
    acc.reverse()
    return "\n".join(acc)

def _ensure_thread(state_key: str) -> List[Dict[str, Any]]:
    """
    Ensure st.session_state[state_key] is a list[dict] with keys: role in {"user","assistant"}, content, ts.
    Migrates old tuple-based history [('student'|'tutor', 'msg')] if present.
    Returns the normalized list.
    """
    st.session_state.setdefault(state_key, [])
    thread = st.session_state[state_key]
    if thread and isinstance(thread[0], tuple):
        # migrate tuples -> dicts; keep a best-effort role mapping
        mapped: List[Dict[str, Any]] = []
        for role, msg in thread:
            r = "user" if role in ("user", "student") else "assistant"
            mapped.append({"role": r, "content": msg, "ts": time.time()})
        st.session_state[state_key] = mapped
        thread = mapped
    return thread

def render_conversation(
    *,
    state_key: str,
    title: str,
    placeholder: str,
    on_ask: Callable[[str, List[Dict[str, Any]]], str],
    clear_label: str = "🗑️ Clear chat",
    before_input: Callable[[], None] | None = None,
) -> None:
    """
    Generic chat renderer:
    - Keeps a thread in st.session_state[state_key]
    - Renders bubbles with st.chat_message + st.chat_input
    - Calls on_ask(user_msg, history_without_current) -> assistant_reply
    - Appends both turns back to the thread
    - Guarantees that the input box appears at the bottom after the latest answer
    """
    st.subheader(title)
    thread = _ensure_thread(state_key)

    # Optional top controls (e.g., info boxes)
    if before_input:
        before_input()

    # --- Transcript (messages) ---
    for msg in thread:
        with st.chat_message("user" if msg["role"] == "user" else "assistant"):
            st.markdown(msg["content"])

    # --- Bottom-placed clear button (appears only when there is history) ---
    if thread:
        spacer, btn_col = st.columns([4, 1])
        with btn_col:
            if st.button(clear_label, key=f"clear_bottom__{state_key}"):
                st.session_state[state_key] = []
                st.rerun()

    # --- Input (unique key per thread avoids collisions when used in several places) ---
    user_q = st.chat_input(placeholder, key=f"chat_input__{state_key}")
    if not user_q:
        return

    # Show user turn immediately
    user_turn = {"role": "user", "content": user_q, "ts": time.time()}
    thread.append(user_turn)
    with st.chat_message("user"):
        st.markdown(user_q)

    # Call the app-provided handler with history (excluding the new turn)
    prior = thread[:-1]
    with st.chat_message("assistant"):
        with st.spinner("Thinking…"):
            answer = on_ask(user_q, prior)
            st.markdown(answer)

    # Append assistant turn
    thread.append({"role": "assistant", "content": answer, "ts": time.time()})

    # Re-run so Streamlit re-renders the full transcript and places a fresh input box
    # *after* the last answer (so the input is always at the bottom).
    st.rerun()

def render_footer():
    # Simple, non-sticky footer at the end of the page
    st.divider()
    st.caption(
        "[AI & Privacy Notice](?show_privacy=1) · "
        "© 2026 Stephan Balthasar · This app uses AI & LLMs; outputs may be inaccurate; no liability. "
        "Feedback is not a grade predictor."
    )
    
# === PATCH 1: load the notice (Markdown file) ===
def load_privacy_notice():
    file_path = os.path.join("assets", "Notice.md")
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return "# Privacy Notice\n\nThe notice could not be loaded. Please contact the administrator."

# === PATCH 1: privacy overlay (opens via ?show_privacy=1) ===
def _get_query_params():
    try:
        return st.experimental_get_query_params()  # legacy API
    except Exception:
        try:
            return dict(st.query_params)  # newer API
        except Exception:
            return {}

def render_privacy_overlay_if_requested():
    qp = _get_query_params()
    val = str(qp.get("show_privacy", ["0"])[0]).lower()
    show = val in ("1", "true", "yes", "y", "on")
    if not show:
        return

    notice_md = load_privacy_notice()
    st.title("AI & Privacy Notice")
    st.markdown(notice_md)
    st.markdown("[← Back to the app](?)")
    render_footer()
    st.stop()

# --- minimalist logger: uses only LOG_GIST_TOKEN + GIST_ID ---
def update_gist(new_entry):
    """
    Append [timestamp, event, role] to b2_log.csv in a GitHub Gist.
    Uses a dedicated token only: st.secrets['LOG_GIST_TOKEN'].
    If not configured, this function silently no-ops.
    """
    token = st.secrets.get("LOG_GIST_TOKEN")
    gist_id = st.secrets.get("GIST_ID")
    if not token or not gist_id:
        return  # no-op if logging is not configured

    url = f"https://api.github.com/gists/{gist_id}"
    headers = {"Authorization": f"token {token}"}

    # 1) Read current b2_log.csv (or start with the header)
    try:
        r = requests.get(url, headers=headers, timeout=8)
        files = r.json().get("files", {}) if r.status_code == 200 else {}
        content = files.get("b2_log.csv", {}).get("content", "")
        lines = [ln for ln in content.splitlines() if ln.strip()] or ["timestamp,event,role"]
    except Exception:
        lines = ["timestamp,event,role"]

    # 2) Append the new entry and push
    lines.append(",".join(new_entry))
    payload = {"files": {"b2_log.csv": {"content": "\n".join(lines)}}}
    try:
        requests.patch(url, headers=headers, data=json.dumps(payload), timeout=8)
    except Exception:
        # Best-effort logging — ignore network/api errors to keep UX smooth
        pass

# --- Load booklet index (server-side; users never see this file) ---
from app.bootstrap_booklet import load_booklet_index
INDEX = load_booklet_index()  # {"paragraphs": [...], "chapters": [...]}

from app.bootstrap_cases import load_cases
CASES = load_cases()

from mentor.rag.booklet_retriever import ParagraphRetriever
from mentor.engines.chat_engine import ChatEngine
from mentor.engines.feedback_engine import FeedbackEngine
from mentor.llm.groq import GroqClient

# Global width cap for a professional look (applies to all pages)
st.markdown("""
<style>
  /* Constrain the main content and center it */
  .block-container, section.main > div {
    max-width: 1120px !important;   /* pick 1040–1160px if you prefer */
    margin: 0 auto !important;
  }
</style>
""", unsafe_allow_html=True)

# === Load brand CSS (scoped details not covered by theme) ===
try:
    with open("assets/theme.css", "r", encoding="utf-8") as _f:
        st.markdown(f"<style>{_f.read()}</style>", unsafe_allow_html=True)
except Exception:
    pass

# === PATCH 2: always-on footer and optional overlay ===
render_privacy_overlay_if_requested()

# === PATCH 3: session flags ===
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "role" not in st.session_state:
    st.session_state.role = None

# === PATCH 3: login gate ===
if not st.session_state.authenticated:
    # Hide the sidebar on the landing page only
    st.markdown(
        """
        <style>
        div[data-testid="stSidebar"] { display: none !important; }
        /* Slightly tighten top/bottom padding while the sidebar is hidden */
        .block-container { padding-top: 0.75rem !important; }
        </style>
        """,
        unsafe_allow_html=True,
    )
    
    # Landing
    render_brand_bar_aligned(
        icon_src="assets/b2_logo_1024.png",
        title="B's Bot",
        subhead="Your AI Mentor for EU Capital Markets Law.",
        bar_height_desktop=44, bar_height_mobile=38,
        logo_top_nudge_px=0, title_nudge_px=3, sub_nudge_px=-3
    )
           
    STUDENT_PIN = st.secrets.get("STUDENT_PIN")
    TUTOR_PIN   = st.secrets.get("TUTOR_PIN")

    # --- Simple, robust login form ---
    # Requirements: show password + checkbox at the same time; allow Enter OR button click.
    with st.form(key="login_form", clear_on_submit=False):
        pin = st.text_input("Enter password", type="password")
        agree = st.checkbox(
            "I have read the AI & Privacy Notice (see link below) and will not include personal data in my submissions. "
        )
    
        submitted = st.form_submit_button("Continue", type="primary")
    
        # --- Inline footer, rendered inside the form on the landing page ---
        st.caption(
            "[AI & Privacy Notice](?show_privacy=1) · "
            "© 2026 Stephan Balthasar · This app uses AI & LLMs; outputs may be inaccurate; no liability. "
            "Feedback is not a grade predictor."
        )
    
        if submitted:
            role_detected = None
            if pin and STUDENT_PIN and pin == STUDENT_PIN:
                role_detected = "student"
            elif pin and TUTOR_PIN and pin == TUTOR_PIN:
                role_detected = "tutor"
    
            messages = []
            if role_detected is None:
                messages.append("Incorrect password")
            if not agree:
                messages.append("Tick the box first")
    
            if messages:
                for m in messages:
                    st.error(m)
            else:
                st.session_state.authenticated = True
                st.session_state.role = role_detected
                if role_detected == "student":
                    update_gist([time.strftime("%Y-%m-%d %H:%M:%S"), "LOGIN", "student"])
                st.rerun()
    st.stop()

# Compact brand bar (authenticated pages only)
render_brand_bar_aligned(
    icon_src="assets/b2_logo_1024.png",
    title="B's Bot",
    subhead="Your AI Mentor for EU Capital Markets Law.",
    bar_height_desktop=44, bar_height_mobile=38,
    logo_top_nudge_px=0, title_nudge_px=3, sub_nudge_px=-3
)

# --- Build retrievers once ---
para_retriever = ParagraphRetriever(INDEX["paragraphs"])

# --- LLM client ---
llm_api_key = st.secrets.get("GROQ_API_KEY")
if not llm_api_key:
    st.error("Missing GROQ_API_KEY in secrets.")
    st.stop()
llm = GroqClient(api_key=llm_api_key)

# --- Engines ---
chat_engine = ChatEngine(
    llm=llm,
    booklet_index=INDEX,          # kept for legacy codepaths / counts elsewhere
    booklet_retriever=para_retriever,
    web_retriever=None
)
feedback_engine = FeedbackEngine(llm=llm)

# --- Sidebar controls ---
with st.sidebar:
    st.caption(f"📖 Booklet loaded — {len(INDEX['chapters'])} chapters / {len(INDEX['paragraphs'])} paragraphs")

    model = st.selectbox(
        "Model",
        ["llama-3.1-8b-instant", "llama-3.3-70b-versatile"],
        index=0,
        help=(
            "Model choice:\n"
            "• llama‑3.1‑8b‑instant → faster, cheaper; good for drafts and everyday Q&A.\n"
            "• llama‑3.3‑70b‑versatile → slower, more capable; better for nuanced legal analysis."
        ),
    )

    temp = st.slider(
        "Temperature",
        0.0, 1.0, 0.2, 0.05,
        help=(
            "Controls randomness.\n"
            "• 0.0–0.3 → deterministic, best for legal reasoning.\n"
            "• 0.4–0.7 → more creative.\n"
            "Higher values may produce inconsistent answers."
        ),
    )

    if st.button("Reload booklet index (server cache)"):
        st.cache_data.clear()
        st.success("Re-loaded. Re-run the action to use the latest JSON.")

    st.divider()
    debug_signals = st.toggle("🔧 Show signal debugger", value=False, help=(
        "Show the raw signals detected by extract_signals() for the last query "
        "(type, surface, canonical, confidence, expanded)."
    ))
    st.session_state["_show_signal_debugger"] = debug_signals

    if debug_signals:
        st.markdown("#### Signal Debugger")
        last = st.session_state.get("_last_signals")
        if not last:
            st.caption("No signals yet — ask a question to see them here.")
        else:
            # Render a compact table
            import pandas as pd
            df = pd.DataFrame(last)
            # Nice, narrow columns for readability
            st.dataframe(
                df[["type", "surface", "canonical", "confidence", "expanded_preview"]],
                hide_index=True,
                use_container_width=True,
            )
            last_dec = st.session_state.get("_last_router_decision")
            if last_dec:
                st.caption(
                    f"Router → {last_dec.get('label')} · v={last_dec.get('v')}"
                )
            cq = st.session_state.get("_last_combined_query")
            if cq:
                with st.expander("Combined query used for RAG"):
                    st.code(cq, language="text")
            # --- Pinned case badge (optional UI) ---
            fc = st.session_state.get("_focus_case")
            if fc and (fc.get("canonical") or "").strip():
                label = fc.get("display") or fc.get("canonical")
                if fc.get("docket"):
                    label = f"{label} ({fc['docket']})"
                st.markdown(f"**📌 Pinned case:** {label}")
                if st.button("Unpin case"):
                    st.session_state["_focus_case"] = None

            
# --- Tabs: Feedback + Tutor chat ---
tab_feedback, tab_chat = st.tabs(["📝 Exam Assistant", "💬 EUCapML Booklet"])

# Small helper: persist latest run per case+question
def _key(case_id: str, q_label: str) -> str:
    return f"{case_id}::{q_label}"

# --- REPLACEMENT FEEDBACK TAB (full block) ----

with tab_feedback:

    st.subheader("Exam Assistant")

    # -----------------------------
    # 1. CASE SELECTION (kept unchanged)
    # -----------------------------
    case_titles = [c.get("title", c.get("id", "Untitled case")) for c in CASES]
    sel_case_title = st.selectbox("Select case", case_titles, index=0)
    sel_case = next(c for c in CASES if c.get("title", c.get("id")) == sel_case_title)
    sel_case_id = sel_case.get("id", "unknown")

    # -----------------------------
    # 2. CASE DESCRIPTION
    # -----------------------------
    st.markdown("### Case description")
    st.write(sel_case.get("description", "—"))

    # -----------------------------
    # 3. QUESTION SELECTION
    # -----------------------------
    q_count = int(sel_case.get("question_count", 1))
    q_labels = [f"Question {i+1}" for i in range(max(1, q_count))]
    q_label = st.selectbox("Which question are you working on?", q_labels, index=0)
    q_index = q_labels.index(q_label)

    st.divider()
    
    # -----------------------------
    # 4. WORKFLOW CHOICE
    # -----------------------------
    workflow = st.radio(
        "Choose your workflow:",
        ["Help me prepare an answer", "I have an answer ready to submit"],
        horizontal=False
    )

    st.divider()


    # -----------------------------
    # 4. PLAN WORKFLOW
    # -----------------------------
    if workflow == "Help me prepare an answer":
        st.markdown("## Plan your answer")
        st.markdown("The app will help you build a structured outline based on the case and model solution.")

        # Load model answer slice
        sections = sel_case.get("model_answer_sections") or []
        model_slice = sections[q_index] if (0 <= q_index < len(sections)) else ""

        if st.button("Generate plan", type="primary"):
            if st.session_state.get("role") == "student":
                update_gist([time.strftime("%Y-%m-%d %H:%M:%S"), "PLAN", "student"])
            with st.spinner("Thinking..."):
                plan = feedback_engine.plan_answer(
                    case_text=sel_case.get("description", ""),
                    question=q_label,
                    model_answer_slice=model_slice,
                    booklet_text="",       # your design: no booklet grounding in plan
                    model=model,
                    temperature=temp
                )
            st.session_state["plan_output"] = plan

        # Display generated plan
        if "plan_output" in st.session_state:
            st.markdown("### Suggested solution structure")
            st.markdown(st.session_state["plan_output"])


    # -----------------------------
    # 5. EVALUATE WORKFLOW
    # -----------------------------
    if workflow == "I have an answer ready to submit":

        st.markdown("## Submit your exam answer")

        # text input
        answer = st.text_area(
            "Your answer",
            height=240,
            key=f"answer::{sel_case_id}::{q_label}"
        )

        # on evaluate
        if st.button("Evaluate my answer", type="primary"):
            if st.session_state.get("role") == "student":
                update_gist([time.strftime("%Y-%m-%d %H:%M:%S"), "EVALUATE", "student"])
            sections = sel_case.get("model_answer_sections") or []
            auto_slice = sections[q_index] if (0 <= q_index < len(sections)) else None
            effective_model = (auto_slice or "").strip()

            if not effective_model or not answer.strip():
                st.warning("Missing model answer slice or student answer.")
            else:
                with st.spinner("Evaluating..."):
                    fb = feedback_engine.evaluate_answer(
                        student_answer=answer,
                        model_answer=effective_model,
                        model=model,
                        temperature=temp
                    )

                # persist all relevant state
                st.session_state["exam_answer"] = answer
                st.session_state["exam_feedback"] = fb
                st.session_state["chat_history"] = []    # reset chat thread

        # SHOW RESULTS AFTER EVALUATION
        if "exam_feedback" in st.session_state:

            st.markdown("## Your submitted answer")
            st.markdown(st.session_state["exam_answer"])

            st.markdown("## Structured feedback")
            st.markdown(st.session_state["exam_feedback"])

            # -----------------------------
            # DOCX DOWNLOAD
            # -----------------------------
            from docx import Document
            from io import BytesIO

            def make_docx():
                doc = Document()
                doc.add_heading(f"Feedback – {sel_case_title} – {q_label}", level=1)

                doc.add_heading("Student Answer", level=2)
                doc.add_paragraph(st.session_state["exam_answer"])

                doc.add_heading("Feedback", level=2)
                doc.add_paragraph(st.session_state["exam_feedback"])

                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                return buf

            st.download_button(
                "📄 Download feedback (.docx)",
                data=make_docx(),
                file_name=f"feedback_{sel_case_id}_{q_label.replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            st.divider()

            # -----------------------------
            # 6. FOLLOW-UP CHAT
            # -----------------------------
            # --- Follow-up discussion (reuses the same conversation component) ---
            with st.container():
                def on_ask_followup(user_q: str, history: List[Dict[str, Any]]) -> str:
                    # Build the history in the legacy tuple format that your feedback_engine expects
                    tuple_hist = [
                        ("student" if m["role"] == "user" else "tutor", m["content"])
                        for m in history
                    ]
                    context = {
                        "student_answer": st.session_state.get("exam_answer", ""),
                        "feedback": st.session_state.get("exam_feedback", ""),
                        "history": tuple_hist,
                    }
                    # Keep your audit behavior: log only for students
                    if st.session_state.get("role") == "student":
                        update_gist([time.strftime("%Y-%m-%d %H:%M:%S"), "FOLLOW_UP", "student"])
            
                    return feedback_engine.follow_up_with_history(
                        question=user_q,
                        context=context,
                        model=model,
                        temperature=temp,
                    )
            
                # Use the same state key you already reset after evaluation to avoid surprises
                render_conversation(
                    state_key="chat_history",  # we keep the same key to preserve behavior / resets
                    title="Follow-up discussion",
                    placeholder="Add your follow-up question…",
                    on_ask=on_ask_followup,
                    clear_label="🗑️ Clear follow-up thread",
                )

# --- Tutor chat (separate, uncluttered) ---
with tab_chat:
    def build_combined_query(recent_msgs, current_msg, window=3):
        msgs = recent_msgs[-(window - 1):] + [current_msg]
        # Clean combination: just stitch the raw user messages
        return " ".join(msgs)

    # --- Case pinning helpers (for robust follow-ups like "this decision") ---

    def _maybe_pin_case_from_signals(signals):
        """
        If the current RAG turn detected a strong case_name signal, pin it in session state.
        'signals' is the cleaned list you already store in _last_signals for the debugger.
        We select the highest-confidence case_name with confidence >= 0.9.
        """
        if not signals:
            return
        best = None
        for s in signals:
            if (s.get("type") == "case_name"):
                try:
                    conf = float(s.get("confidence", 0.0) or 0.0)
                except Exception:
                    conf = 0.0
                if conf >= 0.9 and (best is None or conf > float(best.get("confidence", 0.0) or 0.0)):
                    best = s
    
        if not best:
            return
    
        canonical = (best.get("canonical") or "").strip()
        if not canonical:
            return
    
        # Try to find a docket-like alias from the expanded set (e.g., "C-45/08")
        expanded = best.get("expanded_preview") or ""
        # Note: expanded_preview is a comma-joined short preview in your sidebar;
        # if you prefer full expanded set, adjust to use the raw signals before cleaning.
        docket = None
        for piece in (expanded.split(",") if expanded else []):
            p = piece.strip()
            if p.startswith(("C-", "T-", "ECLI:", "Joined Cases")):
                docket = p
                break
    
        st.session_state["_focus_case"] = {
            "canonical": canonical,               # lower-cased canonical (from signals)
            "display": canonical,                 # you can map to a nicer title later
            "docket": docket,                     # optional
            "confidence": float(best.get("confidence", 0.0) or 0.0),
            "ts": time.time(),
        }
    
    
    def _augment_with_pinned_case(combined_query: str) -> str:
        """
        Prepend the pinned case (if any) to the combined query so the retriever 'sees'
        the anchor case even when the user asks anaphorically ("this decision").
        """
        fc = st.session_state.get("_focus_case")
        if not fc or not (fc.get("canonical") or "").strip():
            return combined_query
    
        tag = (fc.get("display") or fc.get("canonical") or "").strip()
        if not tag:
            return combined_query
    
        if fc.get("docket"):
            tag = f"{tag} ({fc['docket']})"
    
        # Keep this simple—no labels, just plain natural text before the user's query.
        return f"{tag}. {combined_query}"
    
    def on_ask_tutor(user_q: str, history: List[Dict[str, Any]]) -> str:
        # Optional: keep your student usage ping
        if st.session_state.get("role") == "student":
            update_gist([time.strftime("%Y-%m-%d %H:%M:%S"), "CHAT", "student"])
        
        # --- NEW: capture signals for the sidebar debugger ---
        try:
            sigs = extract_signals(
                user_q,
                gaz=para_retriever.gaz,            # your ParagraphRetriever exposes gaz
                corpus_auto_alias=para_retriever.alias_bi,  # merged alias graph
            )
            # Normalize to a printable structure (don’t mutate original dicts)
            cleaned = []
            for s in (sigs or []):
                cleaned.append({
                    "type": s.get("type"),
                    "surface": s.get("surface"),
                    "canonical": s.get("canonical"),
                    "confidence": round(float(s.get("confidence", 0.0)), 3),
                    # Keep a short preview of the expanded set for readability
                    "expanded_preview": ", ".join(sorted(list(s.get("expanded", set())))[:6]),
                })
            st.session_state["_last_signals"] = cleaned
            # After you computed 'cleaned' signals for the debugger (st.session_state["_last_signals"]):
            try:
                cleaned_signals = st.session_state.get("_last_signals") or []
                _maybe_pin_case_from_signals(cleaned_signals)
            except Exception:
                # Keep UX resilient: never break the chat due to pinning
                pass
        except Exception as e:
            # Keep UX resilient; store the error so you can see it in the panel
            st.session_state["_last_signals"] = [{"type": "ERROR", "surface": "", "canonical": str(e), "confidence": 0.0, "expanded_preview": ""}]

        # Heuristic router (no LLM): counts gazetteer hits (exact or fuzzy)
        # --- Router decision ---
        # Build list of last user messages (not including current)
        recent_user_msgs = [
            m["content"] for m in history if m["role"] == "user"
        ]
        
        # Pass to router
        decision = route(
            user_query=user_q,
            recent_user_messages=recent_user_msgs
        )
        
        # Store for debugger if needed
        st.session_state["_last_router_decision"] = {
            "label": decision["ui_label"],
            "v": decision["router_version"],
        }
        
        # Store router information for the sidebar debugger
        st.session_state["_last_router_decision"] = {
            "mode": decision.get("mode"),
            "conf": decision.get("total_conf"),
            "label": decision.get("ui_label"),
            "v": decision.get("router_version"),
        }
        
        # --- RAG vs Chat output WITHOUT inserting router metadata into the chat ---
        if decision["mode"] == "rag":

            # 1) Build the plain combined query (you already have this)
            combined_q = build_combined_query(recent_user_msgs, user_q, window=3)
            
            # 2) (Optional but recommended) Store the plain version for debugging
            st.session_state["_last_combined_query_plain"] = combined_q
            
            # 3) Augment with the pinned case, if any
            combined_q_aug = _augment_with_pinned_case(combined_q)
            
            # 4) Store the augmented version too (so you can compare in the Signal Debugger)
            st.session_state["_last_combined_query_aug"] = combined_q_aug
            
            # 5) Call RAG with the AUGMENTED combined query
            answer = chat_engine.answer(
                user_query=combined_q_aug,
                model=model,
                temperature=temp,
                max_tokens=700
            )
            return answer
        else:
            answer = chat_engine.assist(
                user_query=user_q,
                model="llama-3.1-8b-instant",
                temperature=0.6,
                max_tokens=350,
            )
            return answer
        
    render_conversation(
        state_key="tutor_chat",
        title="General Chat (Course Material)",
        placeholder="Ask the tutor…",
        on_ask=on_ask_tutor,
        clear_label="🗑️ Clear chat for new topics! ",
    )

# --- Page footer (authenticated pages only) ---
render_footer()
