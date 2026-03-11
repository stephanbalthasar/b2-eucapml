# mentor/booklet/index.py
# Minimal index builder for the EUCapML Booklet
# - paragraphs: every paragraph styled as "Standard with para numbering" → para_num 1..N
# - chapters: every "Heading 1" → chapter_num 1..N, body = text until next Heading 1

from docx import Document

def build_booklet_index_from_docx(docx_path: str) -> dict:
    """
    Returns:
      {
        "paragraphs": [
          { "para_num": 1, "text": "...", "chapter_num": 1, "chapter_title": "Introduction" },
          ...
        ],
        "chapters": [
          { "chapter_num": 1, "title": "Introduction", "text": "…" },
          ...
        ]
      }
    """
    doc = Document(docx_path)

    # 1) Build chapters from Heading 1
    chapters = []
    current_title = None
    current_num = 0
    current_buf = []

    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style == "Heading 1":
            # flush previous chapter
            if current_title is not None:
                chapters.append({
                    "chapter_num": current_num,
                    "title": current_title,
                    "text": "\n".join(current_buf).strip(),
                })
            # start new chapter
            current_num += 1
            current_title = text or f"Chapter {current_num}"
            current_buf = []
        else:
            if current_title is not None and text:
                current_buf.append(text)

    # flush last chapter
    if current_title is not None:
        chapters.append({
            "chapter_num": current_num,
            "title": current_title,
            "text": "\n".join(current_buf).strip(),
        })

    # Build a quick index: map paragraph object id to chapter_num so we can tag paragraphs
    # Because we rebuilt chapter texts from linear order, we make a second pass to tag paras with chapter.
    # Simple heuristic: we walk again through the doc, track the current chapter_num.
    para_items = []
    para_counter = 0
    chapter_cursor = 0

    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()

        if style == "Heading 1":
            chapter_cursor += 1
            continue

        if style == "Standard with para numbering" and text:
            para_counter += 1
            para_items.append({
                "para_num": para_counter,
                "text": text,
                "chapter_num": chapter_cursor if chapter_cursor > 0 else None,
                "chapter_title": chapters[chapter_cursor - 1]["title"] if 0 < chapter_cursor <= len(chapters) else None
            })

    return {
        "paragraphs": para_items,
        "chapters": chapters,
    }


# Optional: simple JSON persistence (run once per new booklet)
def save_index_json(index: dict, json_path: str) -> None:
    import json
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(index, f, ensure_ascii=False, indent=2)


def load_index_json(json_path: str) -> dict:
    import json
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)
